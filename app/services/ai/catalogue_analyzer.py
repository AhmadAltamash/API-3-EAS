import json
import re
import time
from io import BytesIO

from google import genai
from config import Config


class CatalogueAnalyzer:
    """
    Reads an uploaded product catalogue (PDF, DOCX, or plain text) and
    asks Gemini to identify what's being sold, then generates search
    queries aimed at finding US/Canada BUYERS of those products -
    importers, wholesalers, distributors, retailers - rather than
    competitors who make the same thing.
    """

    MODEL = "gemini-3.5-flash"

    MAX_RETRIES = 3

    RETRYABLE_MARKERS = (
        "503", "UNAVAILABLE", "429", "RESOURCE_EXHAUSTED", "overloaded"
    )

    # Keep the prompt (and cost/latency) bounded - a catalogue's opening
    # pages are normally enough to identify the product mix.
    MAX_CHARS = 15000

    def __init__(self):

        self.client = genai.Client(
            api_key=Config.GEMINI_API_KEY
        )

    # -------------------------------------------------------------
    # File text extraction
    # -------------------------------------------------------------

    def extract_text(self, file_storage):
        """
        file_storage: a Werkzeug FileStorage from request.files.
        Returns extracted plain text, or raises ValueError with a
        user-facing message on an unsupported/unreadable file.
        """

        filename = (file_storage.filename or "").lower()

        data = file_storage.read()

        if not data:
            raise ValueError("The uploaded file appears to be empty.")

        if filename.endswith(".pdf"):
            return self._extract_pdf(data)

        if filename.endswith(".docx"):
            return self._extract_docx(data)

        if filename.endswith(".txt") or filename.endswith(".csv"):
            return data.decode("utf-8", errors="ignore")

        raise ValueError(
            "Unsupported file type. Please upload a PDF, DOCX, or TXT catalogue."
        )

    def _extract_pdf(self, data):

        try:
            from pypdf import PdfReader
        except ImportError:
            raise ValueError(
                "PDF support isn't installed yet. Run: pip install pypdf "
                "(see requirements.txt)."
            )

        reader = PdfReader(BytesIO(data))

        text_parts = []

        for page in reader.pages:

            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                continue

        text = "\n".join(text_parts).strip()

        if not text:
            raise ValueError(
                "Couldn't read any text from this PDF - it may be a "
                "scanned/image-only catalogue, which isn't supported yet. "
                "Try a text-based PDF, DOCX, or TXT file instead."
            )

        return text

    def _extract_docx(self, data):

        try:
            from docx import Document
        except ImportError:
            raise ValueError(
                "DOCX support isn't installed yet. Run: pip install "
                "python-docx (see requirements.txt)."
            )

        document = Document(BytesIO(data))

        text = "\n".join(
            p.text for p in document.paragraphs if p.text.strip()
        )

        # Catalogues often list products in tables rather than paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()

        text = text.strip()

        if not text:
            raise ValueError("Couldn't find any readable text in this DOCX file.")

        return text

    # -------------------------------------------------------------
    # AI analysis
    # -------------------------------------------------------------

    def analyze(self, catalogue_text):
        """
        Returns a dict:
        {
            "products": [...],
            "categories": [...],
            "materials": [...],
            "industries": [...],
            "buyer_types": [...],
            "keywords": [{"query": "...", "country": "...", "reason": "..."}]
        }
        """

        trimmed = catalogue_text[: self.MAX_CHARS]

        prompt = f"""
        You are a B2B export sales analyst.

        Read the product catalogue text below and identify what this
        company manufactures/sells.

        CATALOGUE TEXT:
        \"\"\"
        {trimmed}
        \"\"\"

        Based on this, generate search queries designed to find companies
        in the UNITED STATES and CANADA that would realistically BUY these
        products - importers, wholesalers, distributors, and retailers who
        would stock or resell this kind of product.

        Rules:
        - Do NOT suggest queries that would surface other manufacturers of
          the same product (we do not want competitors).
        - Do NOT suggest generic single-word product searches.
        - Each query should combine a buyer type + product category +
          country, e.g. "home decor importer USA" or
          "giftware distributor Canada".
        - Cover a mix of both United States and Canada queries.

        Respond with ONLY valid JSON (no markdown fences, no commentary,
        no trailing commas), exactly matching this structure:

        {{
          "products": ["...", "..."],
          "categories": ["...", "..."],
          "materials": ["...", "..."],
          "industries": ["...", "..."],
          "buyer_types": ["...", "..."],
          "keywords": [
            {{"query": "...", "country": "United States", "reason": "..."}},
            {{"query": "...", "country": "Canada", "reason": "..."}}
          ]
        }}

        Generate between 8 and 14 keyword entries total.
        """

        raw_text = self._generate_with_retry(prompt)

        data = self._parse_json(raw_text)

        data.setdefault("products", [])
        data.setdefault("categories", [])
        data.setdefault("materials", [])
        data.setdefault("industries", [])
        data.setdefault("buyer_types", [])
        data.setdefault("keywords", [])

        return data

    def _generate_with_retry(self, prompt):

        last_error = None

        for attempt in range(1, self.MAX_RETRIES + 1):

            try:

                response = self.client.models.generate_content(
                    model=self.MODEL,
                    contents=prompt
                )

                return response.text or ""

            except Exception as e:

                last_error = e

                error_text = str(e)

                is_retryable = any(
                    marker in error_text
                    for marker in self.RETRYABLE_MARKERS
                )

                print(f"Catalogue AI error (attempt {attempt}/{self.MAX_RETRIES}):", e)

                if not is_retryable or attempt == self.MAX_RETRIES:
                    break

                time.sleep(2 ** attempt)

        raise ValueError(
            f"The AI service is temporarily unavailable. Please try again "
            f"in a moment. ({last_error})"
        )

    def _parse_json(self, raw_text):

        cleaned = raw_text.strip()

        # Strip markdown code fences if Gemini wraps the JSON in them
        cleaned = re.sub(r"^```(json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:

            # Sometimes the model adds a stray sentence before/after the
            # JSON object - try to salvage just the {...} block.
            match = re.search(r"\{.*\}", cleaned, re.DOTALL)

            if match:
                try:
                    return json.loads(match.group(0))
                except json.JSONDecodeError:
                    pass

            raise ValueError(
                "The AI response couldn't be read as valid data. Please try again."
            )
